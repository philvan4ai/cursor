#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成垂钓管理平台需求规格说明书 Word 文档。"""

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


def set_run_font(run, name="宋体", size=12, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_heading_cn(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        if level == 1:
            set_run_font(run, "黑体", 16, True, RGBColor(0x1A, 0x56, 0x3A))
        elif level == 2:
            set_run_font(run, "黑体", 14, True, RGBColor(0x2D, 0x6A, 0x4F))
        else:
            set_run_font(run, "黑体", 12, True, RGBColor(0x40, 0x40, 0x40))
    return h


def add_para(doc, text, size=11, bold=False, first_line=True, space_after=8):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if first_line:
        pf.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    set_run_font(run, "宋体", size, bold)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    pf = p.paragraph_format
    pf.space_after = Pt(4)
    pf.left_indent = Cm(0.74 + level * 0.5)
    run = p.add_run(text)
    set_run_font(run, "宋体", 11)
    return p


def set_cell_text(cell, text, bold=False, size=10, align=WD_ALIGN_PARAGRAPH.LEFT, bg=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    set_run_font(run, "宋体", size, bold)
    if bg:
        shading = cell._element.get_or_add_tcPr()
        shd = shading.makeelement(
            qn("w:shd"),
            {
                qn("w:val"): "clear",
                qn("w:color"): "auto",
                qn("w:fill"): bg,
            },
        )
        shading.append(shd)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, bg="1B4332")
        for run in table.rows[0].cells[i].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            set_run_font(run, "宋体", 10, True, RGBColor(0xFF, 0xFF, 0xFF))
    for r_idx, row in enumerate(rows):
        bg = "F0F7F4" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row):
            set_cell_text(table.rows[r_idx + 1].cells[c_idx], str(val), size=10, bg=bg)
    if col_widths:
        for row in table.rows:
            for idx, w in enumerate(col_widths):
                row.cells[idx].width = Cm(w)
    doc.add_paragraph()
    return table


def build():
    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # ========== 封面 ==========
    for _ in range(3):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("垂钓管理平台")
    set_run_font(run, "黑体", 28, True, RGBColor(0x1B, 0x43, 0x32))

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("产品需求规格说明书")
    set_run_font(run, "黑体", 22, True, RGBColor(0x2D, 0x6A, 0x4F))

    en = doc.add_paragraph()
    en.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = en.add_run("Fishing Management Platform — Product Requirements Document")
    set_run_font(run, "Arial", 11, False, RGBColor(0x66, 0x66, 0x66))

    for _ in range(2):
        doc.add_paragraph()

    meta_items = [
        ("文档版本", "V1.0"),
        ("文档状态", "初稿完善版"),
        ("产品形态", "多端交互原型 / 管理平台"),
        ("角色体系", "用户端 · 钓场端 · 后台端"),
        ("编写说明", "基于已构建的三角色交互原型完善需求"),
        ("日期", "2026-08-11"),
    ]
    meta_table = doc.add_table(rows=len(meta_items), cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(meta_items):
        set_cell_text(meta_table.rows[i].cells[0], k, bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, bg="E8F5E9")
        set_cell_text(meta_table.rows[i].cells[1], v, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
        meta_table.rows[i].cells[0].width = Cm(4)
        meta_table.rows[i].cells[1].width = Cm(10)

    doc.add_page_break()

    # ========== 修订记录 ==========
    add_heading_cn(doc, "修订记录", 1)
    add_table(
        doc,
        ["版本", "日期", "修订人", "修订说明"],
        [
            ["V0.1", "2026-08", "产品组", "原型功能清单初稿（用户/钓场/后台）"],
            ["V1.0", "2026-08-11", "产品组", "完善为正式需求规格：业务流程、规则、非功能、数据与验收"],
        ],
        [2.5, 3, 3, 8],
    )

    # ========== 1. 文档说明 ==========
    add_heading_cn(doc, "1. 文档说明", 1)
    add_heading_cn(doc, "1.1 编写目的", 2)
    add_para(
        doc,
        "本文档在已构建的垂钓平台交互原型基础上，系统化完善产品需求，明确用户端、钓场端、后台端三类角色的功能边界、业务规则、数据指标与验收标准，作为设计、研发、测试与验收的统一依据。",
    )

    add_heading_cn(doc, "1.2 适用范围", 2)
    add_bullet(doc, "产品经理：需求澄清、优先级排序与版本规划")
    add_bullet(doc, "设计：交互原型落地与视觉规范对齐")
    add_bullet(doc, "研发：功能实现、接口约定与状态机设计")
    add_bullet(doc, "测试：用例编写、验收对照与回归范围")
    add_bullet(doc, "运营/钓场：业务规则确认与上线培训")

    add_heading_cn(doc, "1.3 术语与缩写", 2)
    add_table(
        doc,
        ["术语", "说明"],
        [
            ["用户端", "面向垂钓爱好者的客户端，含地图、购票、打卡、个人中心"],
            ["钓场端", "面向钓场运营人员的端，含区域监控、票务核销、打卡列表、统计"],
            ["后台端", "面向平台运营的管理端，含全局指标、收入趋势与收益明细"],
            ["年票", "按自然年或约定有效期计费的通行证，支持推荐码优惠"],
            ["次票", "单次入场/垂钓凭证，固定单价，无推荐码优惠"],
            ["推荐码", "用于年票优惠的营销码，如 FISH2026"],
            ["打卡", "用户基于 GPS 在指定钓场区域完成位置确认并上报"],
            ["核销", "钓场端扫码确认票务有效并记入核销记录"],
            ["脉冲动画", "钓场端地图上已打卡用户点位的实时高亮动效"],
        ],
        [3.5, 12.5],
    )

    # ========== 2. 产品概述 ==========
    add_heading_cn(doc, "2. 产品概述", 1)
    add_heading_cn(doc, "2.1 产品定位", 2)
    add_para(
        doc,
        "垂钓管理平台是连接垂钓用户、钓场运营与平台管理的一体化数字化系统。平台以「购票—打卡—核销—统计」为主链路，通过地图可视化与实时点位，提升用户入场体验与钓场现场运营效率，并为后台提供可运营的收入与客流洞察。",
    )

    add_heading_cn(doc, "2.2 产品目标", 2)
    add_bullet(doc, "用户侧：便捷找场、购票、打卡，清晰掌握票务与成长权益")
    add_bullet(doc, "钓场侧：实时掌握在场用户、快速核销、掌握当日经营数据")
    add_bullet(doc, "平台侧：全局监控用户与钓场经营，拆分收入结构，支持运营决策")
    add_bullet(doc, "体验侧：顶部角色切换 + 底部导航，支持三角色完整演示与验收")

    add_heading_cn(doc, "2.3 用户角色", 2)
    add_table(
        doc,
        ["角色", "典型用户", "核心诉求", "主要入口"],
        [
            ["垂钓用户", "爱好者 / 会员", "找场、购票、打卡、查权益", "用户端"],
            ["钓场运营", "场长 / 检票员", "看点位、核销、看客流收入", "钓场端"],
            ["平台运营", "运营 / 财务", "看全局指标与收益明细", "后台端"],
        ],
        [3, 3.5, 5.5, 3],
    )

    add_heading_cn(doc, "2.4 原型交互总则", 2)
    add_para(
        doc,
        "点击顶部「用户 / 钓场 / 后台」按钮切换角色；切换后底部导航与页面栈随角色切换。各端内部通过底部导航在功能模块间切换。角色切换应保留演示数据一致性（同一用户购票/打卡后，钓场端与后台端对应指标即时或准实时可见）。",
        first_line=True,
    )

    # ========== 3. 业务主流程 ==========
    add_heading_cn(doc, "3. 业务主流程", 1)
    add_heading_cn(doc, "3.1 端到端主链路", 2)
    add_para(doc, "推荐演示路径如下：", first_line=False)
    add_bullet(doc, "① 用户端地图：浏览已开通 / 维护中钓场，查看规则与鱼种")
    add_bullet(doc, "② 用户端购票：购买年票（可填推荐码 FISH2026 享优惠）或次票")
    add_bullet(doc, "③ 用户端打卡：GPS 选区确认打卡，位置上传后台")
    add_bullet(doc, "④ 钓场端：区域地图脉冲展示该用户；打卡点位列表可见")
    add_bullet(doc, "⑤ 钓场端：扫码核销年票/次票，写入核销记录")
    add_bullet(doc, "⑥ 后台端：今日打卡、今日收入、收益明细同步更新")

    add_heading_cn(doc, "3.2 状态流转（摘要）", 2)
    add_table(
        doc,
        ["对象", "状态", "触发条件", "可见端"],
        [
            ["钓场区域", "已开通 / 维护中", "后台配置或钓场状态变更", "用户地图、后台"],
            ["年票", "未购买 / 有效 / 已过期", "购票成功、到期", "用户「我的」、核销"],
            ["次票", "未购买 / 待使用 / 已核销", "购票、核销", "用户「我的」、核销"],
            ["打卡记录", "已打卡（有效期内）", "GPS 确认成功", "钓场地图/列表、后台"],
            ["核销记录", "已核销", "扫码核销成功", "钓场核销页、后台收入"],
        ],
        [3, 4, 5, 4],
    )

    # ========== 4. 用户端 ==========
    add_heading_cn(doc, "4. 用户端功能需求", 1)
    add_para(
        doc,
        "用户端面向垂钓爱好者，提供地图找场、购票、GPS 打卡与个人权益中心四大模块。",
    )

    add_heading_cn(doc, "4.1 地图模块", 2)
    add_heading_cn(doc, "4.1.1 功能说明", 3)
    add_para(
        doc,
        "在地图上展示已开通与维护中的钓场区域；用户点击区域可查看垂钓规则、鱼种信息等详情，作为购票与打卡的前置信息入口。",
    )
    add_heading_cn(doc, "4.1.2 需求细则", 3)
    add_table(
        doc,
        ["编号", "需求描述", "优先级"],
        [
            ["U-MAP-01", "地图渲染全部可见钓场区域，区分「已开通」「维护中」状态样式", "P0"],
            ["U-MAP-02", "点击区域弹出/进入详情：垂钓规则、开放时段、鱼种信息、注意事项", "P0"],
            ["U-MAP-03", "维护中钓场允许查看信息，但购票/打卡入口置灰或提示不可用", "P1"],
            ["U-MAP-04", "支持定位到用户当前位置，并高亮附近钓场（原型可模拟）", "P1"],
            ["U-MAP-05", "详情页提供「去购票」「去打卡」快捷入口（在已开通时）", "P1"],
        ],
        [3, 11, 2],
    )
    add_heading_cn(doc, "4.1.3 验收要点", 3)
    add_bullet(doc, "已开通与维护中视觉可区分，点击均可查看规则与鱼种")
    add_bullet(doc, "维护中不可完成有效购票/打卡（或明确拦截提示）")

    add_heading_cn(doc, "4.2 购票模块", 2)
    add_heading_cn(doc, "4.2.1 功能说明", 3)
    add_para(
        doc,
        "支持年票与次票两类票务。年票支持推荐码优惠（示例：输入 FISH2026 享受 ¥688 优惠价）；次票单价 ¥68/次，不支持推荐码。",
    )
    add_heading_cn(doc, "4.2.2 票种与价格规则", 3)
    add_table(
        doc,
        ["票种", "标准价", "优惠规则", "有效期/使用", "说明"],
        [
            ["年票", "以产品配置为准（原型展示优惠后 ¥688）", "推荐码 FISH2026 → ¥688", "按年/剩余天数计", "可核销多次（规则可配）"],
            ["次票", "¥68 / 次", "无推荐码优惠", "单次有效，核销后失效", "适合偶发到访"],
        ],
        [2.5, 4, 4, 3.5, 3],
    )
    add_heading_cn(doc, "4.2.3 需求细则", 3)
    add_table(
        doc,
        ["编号", "需求描述", "优先级"],
        [
            ["U-TKT-01", "购票页展示年票、次票卡片：价格、权益摘要、适用说明", "P0"],
            ["U-TKT-02", "年票支持输入推荐码；正确码 FISH2026 时价格变为 ¥688 并明示优惠", "P0"],
            ["U-TKT-03", "推荐码错误/过期时提示失败，价格保持原价", "P0"],
            ["U-TKT-04", "次票固定 ¥68，不展示推荐码输入或输入无效提示", "P0"],
            ["U-TKT-05", "支付成功后生成电子凭证（二维码/条码），同步「我的」票务状态", "P0"],
            ["U-TKT-06", "已持有效年票时，重复购买需二次确认或引导续费策略（可配置）", "P2"],
            ["U-TKT-07", "支付流程可在原型中模拟成功/失败，失败可重试", "P1"],
        ],
        [3, 11, 2],
    )
    add_heading_cn(doc, "4.2.4 业务规则", 3)
    add_bullet(doc, "推荐码仅作用于年票；次票不可叠加推荐码")
    add_bullet(doc, "同一订单仅允许一个有效推荐码")
    add_bullet(doc, "购票成功后，后台端「今日收入」与对应钓场收益明细应增加")
    add_bullet(doc, "凭证需支持钓场端扫码核销")

    add_heading_cn(doc, "4.3 打卡模块", 2)
    add_heading_cn(doc, "4.3.1 功能说明", 3)
    add_para(
        doc,
        "用户通过 GPS 定位选择垂钓区域，确认打卡后将位置上传至后台；钓场端可实时看到该用户点位（地图脉冲 + 列表）。",
    )
    add_heading_cn(doc, "4.3.2 需求细则", 3)
    add_table(
        doc,
        ["编号", "需求描述", "优先级"],
        [
            ["U-CHK-01", "打卡页获取/模拟 GPS，展示可选钓场区域列表或地图选区", "P0"],
            ["U-CHK-02", "确认打卡前校验：区域已开通、用户持有效票（年票或未核销次票）", "P0"],
            ["U-CHK-03", "确认后生成打卡记录：用户、区域、经纬度、时间，并上传后台", "P0"],
            ["U-CHK-04", "打卡成功后提示成功，并引导可在钓场端查看实时点位", "P1"],
            ["U-CHK-05", "定位失败或超出区域范围时给出明确失败原因", "P0"],
            ["U-CHK-06", "支持当日多次打卡策略（可配置：同区冷却时间 / 跨区规则）", "P2"],
        ],
        [3, 11, 2],
    )
    add_heading_cn(doc, "4.3.3 验收要点", 3)
    add_bullet(doc, "打卡成功后，钓场端区域地图出现脉冲点位，列表出现该用户及时间")
    add_bullet(doc, "后台「今日打卡」计数 +1")

    add_heading_cn(doc, "4.4 我的（个人中心）", 2)
    add_heading_cn(doc, "4.4.1 功能说明", 3)
    add_para(
        doc,
        "展示用户资产与成长信息：年票剩余天数、次票状态、积分、等级、每日任务进度。",
    )
    add_heading_cn(doc, "4.4.2 需求细则", 3)
    add_table(
        doc,
        ["编号", "需求描述", "优先级"],
        [
            ["U-ME-01", "展示年票剩余天数；无年票时展示「未开通」与购票引导", "P0"],
            ["U-ME-02", "展示次票状态：无票 / 待使用 / 已核销（及最近核销时间）", "P0"],
            ["U-ME-03", "展示当前积分与等级（等级规则可配置，原型给示例档位）", "P1"],
            ["U-ME-04", "展示每日任务列表与进度（如：完成打卡、完成核销到访等）", "P1"],
            ["U-ME-05", "任务完成后刷新进度与积分；达到阈值可升级", "P1"],
            ["U-ME-06", "提供票务凭证入口，便于出示二维码", "P0"],
        ],
        [3, 11, 2],
    )
    add_heading_cn(doc, "4.4.3 示例成长规则（可配置）", 3)
    add_table(
        doc,
        ["等级", "积分门槛（示例）", "权益示例"],
        [
            ["Lv1 新手钓友", "0", "基础打卡、购票"],
            ["Lv2 熟练钓友", "100", "任务积分加速"],
            ["Lv3 资深钓友", "500", "专属推荐活动曝光"],
        ],
        [4, 5, 7],
    )

    # ========== 5. 钓场端 ==========
    add_heading_cn(doc, "5. 钓场端功能需求", 1)
    add_para(
        doc,
        "钓场端面向单钓场运营，仅展示本钓场数据，强调实时在场态势、票务核销与当日经营概览。",
    )

    add_heading_cn(doc, "5.1 区域地图", 2)
    add_heading_cn(doc, "5.1.1 功能说明", 3)
    add_para(
        doc,
        "仅显示本钓场地图范围，实时展示已打卡用户点位，并以脉冲动画突出在场动态。",
    )
    add_heading_cn(doc, "5.1.2 需求细则", 3)
    add_table(
        doc,
        ["编号", "需求描述", "优先级"],
        [
            ["V-MAP-01", "地图范围限定为本钓场，不展示其他钓场", "P0"],
            ["V-MAP-02", "实时（或轮询）展示当前有效打卡用户点位", "P0"],
            ["V-MAP-03", "点位使用脉冲动画，点击可查看用户脱敏信息与打卡时间", "P0"],
            ["V-MAP-04", "用户离场/打卡失效后点位移除或置灰", "P1"],
            ["V-MAP-05", "支持点位数量角标或在场人数汇总", "P1"],
        ],
        [3, 11, 2],
    )

    add_heading_cn(doc, "5.2 票务核销", 2)
    add_heading_cn(doc, "5.2.1 功能说明", 3)
    add_para(
        doc,
        "支持扫码核销年票与次票，核销成功后展示结果，并保留核销记录列表。",
    )
    add_heading_cn(doc, "5.2.2 需求细则", 3)
    add_table(
        doc,
        ["编号", "需求描述", "优先级"],
        [
            ["V-VFY-01", "支持扫描用户票务二维码完成核销（原型可模拟扫码输入）", "P0"],
            ["V-VFY-02", "年票：核销记入记录，票仍有效（按规则扣减次数或仅登记入场）", "P0"],
            ["V-VFY-03", "次票：核销后状态变为已核销，不可再次使用", "P0"],
            ["V-VFY-04", "无效票/过期票/已核销次票给出明确失败原因", "P0"],
            ["V-VFY-05", "核销记录列表：时间、票种、用户脱敏、结果、操作员（可模拟）", "P0"],
            ["V-VFY-06", "核销成功联动收入统计（若购票已计入收入，则核销侧重客流；若次票核销才确认收入，需在配置中明确）", "P1"],
        ],
        [3, 11, 2],
    )
    add_heading_cn(doc, "5.2.3 推荐确认规则（写入需求）", 3)
    add_para(
        doc,
        "默认建议：购票成功即计入收入；核销用于入场校验与客流统计。若业务要求「次票核销确认收入」，需在配置中切换，并在后台收益明细中保持一致口径。原型演示采用「购票即计入收入」。",
        first_line=True,
    )

    add_heading_cn(doc, "5.3 打卡点位列表", 2)
    add_heading_cn(doc, "5.3.1 功能说明", 3)
    add_para(doc, "以列表形式展示当前垂钓用户及打卡时间，与区域地图点位数据同源。")
    add_heading_cn(doc, "5.3.2 需求细则", 3)
    add_table(
        doc,
        ["编号", "需求描述", "优先级"],
        [
            ["V-LST-01", "列表字段：用户昵称（可脱敏）、打卡时间、区域名、票种标签", "P0"],
            ["V-LST-02", "按打卡时间倒序；支持下拉刷新", "P0"],
            ["V-LST-03", "与地图点位联动：列表选中可在地图高亮", "P2"],
            ["V-LST-04", "空状态提示「暂无在场用户」", "P1"],
        ],
        [3, 11, 2],
    )

    add_heading_cn(doc, "5.4 数据统计", 2)
    add_heading_cn(doc, "5.4.1 功能说明", 3)
    add_para(doc, "展示本钓场经营概览：今日收入、客流、年票/次票持有数。")
    add_heading_cn(doc, "5.4.2 指标定义", 3)
    add_table(
        doc,
        ["指标", "口径", "刷新"],
        [
            ["今日收入", "本钓场当日购票成功金额合计（默认口径）", "准实时"],
            ["客流", "本日有效打卡人数或核销入场人数（需在界面标注口径）", "准实时"],
            ["年票持有数", "当前仍有效且归属/适用于本钓场的年票数量", "定时/事件"],
            ["次票持有数", "已购未核销的次票数量（待使用）", "准实时"],
        ],
        [3.5, 10, 2.5],
    )
    add_heading_cn(doc, "5.4.3 需求细则", 3)
    add_table(
        doc,
        ["编号", "需求描述", "优先级"],
        [
            ["V-STA-01", "卡片展示今日收入、客流、年票持有数、次票持有数", "P0"],
            ["V-STA-02", "指标口径在页面有简短说明或问号提示", "P1"],
            ["V-STA-03", "与用户购票/打卡/核销操作联动更新（演示可即时）", "P0"],
        ],
        [3, 11, 2],
    )

    # ========== 6. 后台端 ==========
    add_heading_cn(doc, "6. 后台端功能需求", 1)
    add_para(
        doc,
        "后台端面向平台运营，提供跨钓场的核心指标、收入趋势与收益明细拆分。",
    )

    add_heading_cn(doc, "6.1 核心指标", 2)
    add_table(
        doc,
        ["指标", "定义", "优先级"],
        [
            ["总用户", "平台累计注册/演示用户数", "P0"],
            ["钓场数", "已接入平台的钓场数量（含维护中）", "P0"],
            ["今日打卡", "自然日内全平台打卡成功次数或去重人数（界面标注）", "P0"],
            ["今日收入", "自然日内全平台购票成功金额合计", "P0"],
        ],
        [3, 11, 2],
    )
    add_table(
        doc,
        ["编号", "需求描述", "优先级"],
        [
            ["A-KPI-01", "顶部/首屏展示四大核心指标卡片", "P0"],
            ["A-KPI-02", "用户打卡、购票后对应指标即时更新（演示环境）", "P0"],
            ["A-KPI-03", "指标支持环比/较昨日差值展示（可选）", "P2"],
        ],
        [3, 11, 2],
    )

    add_heading_cn(doc, "6.2 收入趋势", 2)
    add_para(
        doc,
        "提供周收入柱状图，并展示本月收入与今日收入的对比信息，帮助运营快速判断节奏。",
    )
    add_table(
        doc,
        ["编号", "需求描述", "优先级"],
        [
            ["A-TRD-01", "周收入柱状图：近 7 日每日收入", "P0"],
            ["A-TRD-02", "展示本月累计收入、今日收入及二者对比关系（差值/占比）", "P0"],
            ["A-TRD-03", "图表支持悬停查看具体数值（Web）或点击查看（移动）", "P1"],
            ["A-TRD-04", "无数据日显示 0，避免断轴误导", "P1"],
        ],
        [3, 11, 2],
    )

    add_heading_cn(doc, "6.3 收益明细", 2)
    add_para(
        doc,
        "按钓场拆分年票/次票收益，展示打卡次数，并以收益占比条直观比较各场贡献。",
    )
    add_table(
        doc,
        ["编号", "需求描述", "优先级"],
        [
            ["A-DTL-01", "列表/表格：钓场名称、年票收益、次票收益、打卡次数、收益合计", "P0"],
            ["A-DTL-02", "收益占比条：该钓场收益 / 平台总收益", "P0"],
            ["A-DTL-03", "支持按收益合计排序", "P1"],
            ["A-DTL-04", "年票与次票收益分色或分列，避免口径混淆", "P0"],
            ["A-DTL-05", "点击钓场可下钻至更细记录（可选增强）", "P2"],
        ],
        [3, 11, 2],
    )

    # ========== 7. 信息架构与导航 ==========
    add_heading_cn(doc, "7. 信息架构与导航", 1)
    add_heading_cn(doc, "7.1 全局角色切换", 2)
    add_bullet(doc, "顶部固定：「用户」「钓场」「后台」三按钮互斥切换")
    add_bullet(doc, "切换角色时加载对应底部导航与默认首页")
    add_bullet(doc, "演示数据跨端共享，保证操作可观测")

    add_heading_cn(doc, "7.2 底部导航结构", 2)
    add_table(
        doc,
        ["端", "Tab1", "Tab2", "Tab3", "Tab4"],
        [
            ["用户端", "地图", "购票", "打卡", "我的"],
            ["钓场端", "区域地图", "票务核销", "打卡点位", "数据统计"],
            ["后台端", "核心指标", "收入趋势", "收益明细", "—（可合并为仪表盘多区块）"],
        ],
        [3, 3.5, 3.5, 3.5, 4],
    )
    add_para(
        doc,
        "说明：后台端若采用单页仪表盘，可将核心指标、趋势图、明细纵向排布，仍需保证模块标题与本文档一致，便于验收对照。",
        first_line=True,
    )

    # ========== 8. 非功能需求 ==========
    add_heading_cn(doc, "8. 非功能需求", 1)
    add_heading_cn(doc, "8.1 体验与性能", 2)
    add_bullet(doc, "角色切换与 Tab 切换响应流畅，演示环境无明显卡顿")
    add_bullet(doc, "地图点位脉冲动画流畅，不影响列表滚动")
    add_bullet(doc, "关键操作（购票、打卡、核销）需有明确成功/失败反馈")

    add_heading_cn(doc, "8.2 兼容与适配", 2)
    add_bullet(doc, "原型需同时适配桌面与移动视口基本可用性")
    add_bullet(doc, "文字与按钮可点击区域符合移动端操作习惯")

    add_heading_cn(doc, "8.3 安全与隐私", 2)
    add_bullet(doc, "钓场端、后台展示用户信息默认脱敏（如昵称部分隐藏）")
    add_bullet(doc, "定位数据仅用于打卡校验与在场展示，不做无关用途说明外的扩散")
    add_bullet(doc, "核销与购票操作需可追溯（记录时间与结果）")

    add_heading_cn(doc, "8.4 可用性", 2)
    add_bullet(doc, "空状态、加载中、失败态均有文案")
    add_bullet(doc, "维护中钓场、无效推荐码、定位失败等异常有明确指引")

    # ========== 9. 数据与接口（逻辑） ==========
    add_heading_cn(doc, "9. 数据对象与逻辑接口（摘要）", 1)
    add_heading_cn(doc, "9.1 核心数据对象", 2)
    add_table(
        doc,
        ["对象", "关键字段", "备注"],
        [
            ["User", "id, nickname, points, level, tasks_progress", "用户成长"],
            ["Venue", "id, name, status(open/maintain), rules, fish_species, boundary", "钓场"],
            ["Ticket", "id, user_id, type(year/once), price, promo_code, status, expire_at", "票务"],
            ["CheckIn", "id, user_id, venue_id, lat, lng, time", "打卡"],
            ["VerifyLog", "id, ticket_id, venue_id, time, result, operator", "核销"],
            ["Revenue", "venue_id, date, year_ticket_amount, once_ticket_amount, checkin_count", "收益汇总"],
        ],
        [3, 9, 4],
    )

    add_heading_cn(doc, "9.2 关键逻辑能力（原型/正式实现均可对照）", 2)
    add_bullet(doc, "GET venues / venue detail（规则、鱼种、状态）")
    add_bullet(doc, "POST purchase ticket（含 promo_code 校验）")
    add_bullet(doc, "POST check-in（GPS + 票务校验）")
    add_bullet(doc, "POST verify ticket（扫码核销）")
    add_bullet(doc, "GET venue realtime check-ins")
    add_bullet(doc, "GET admin dashboard metrics / weekly revenue / venue revenue breakdown")

    # ========== 10. 演示账号与样例数据 ==========
    add_heading_cn(doc, "10. 演示数据与操作提示", 1)
    add_heading_cn(doc, "10.1 推荐演示数据", 2)
    add_table(
        doc,
        ["项目", "示例值"],
        [
            ["年票优惠推荐码", "FISH2026 → ¥688"],
            ["次票价格", "¥68 / 次"],
            ["钓场状态", "至少一个「已开通」、一个「维护中」便于对比"],
            ["打卡动画", "钓场端点位脉冲可见"],
        ],
        [5, 11],
    )
    add_heading_cn(doc, "10.2 操作提示（写入产品说明）", 2)
    add_para(
        doc,
        "点击顶部「用户 / 钓场 / 后台」按钮切换角色，底部导航切换各功能页面。建议按「购票 → 打卡 → 钓场查看点位/核销 → 后台看指标」顺序完整体验主链路。",
        first_line=True,
    )

    # ========== 11. 验收标准 ==========
    add_heading_cn(doc, "11. 验收标准（摘要）", 1)
    add_table(
        doc,
        ["编号", "验收项", "通过标准"],
        [
            ["AC-01", "三角色切换", "顶部可切换用户/钓场/后台，页面与导航正确"],
            ["AC-02", "地图状态", "已开通/维护中可区分，详情含规则与鱼种"],
            ["AC-03", "年票优惠", "FISH2026 正确优惠至 ¥688；错误码失败提示"],
            ["AC-04", "次票购买", "¥68 无推荐码优惠，购票后「我的」状态更新"],
            ["AC-05", "打卡联动", "打卡后钓场地图脉冲+列表可见，后台今日打卡增加"],
            ["AC-06", "核销", "年票/次票可核销，记录可查；次票核销后不可复用"],
            ["AC-07", "钓场统计", "今日收入、客流、年票/次票持有数展示正确"],
            ["AC-08", "后台趋势与明细", "周柱状图、本月/今日对比、各场收益拆分与占比条完整"],
        ],
        [2.5, 4, 9.5],
    )

    # ========== 12. 范围与后续 ==========
    add_heading_cn(doc, "12. 本期范围与后续规划", 1)
    add_heading_cn(doc, "12.1 本期（原型/MVP）纳入", 2)
    add_bullet(doc, "三角色完整页面与主链路演示")
    add_bullet(doc, "地图、购票（含推荐码）、打卡、核销、统计与后台看板")
    add_heading_cn(doc, "12.2 明确暂不纳入（可后续迭代）", 2)
    add_bullet(doc, "真实支付渠道对接、退款与发票")
    add_bullet(doc, "复杂会员商城、社区内容、赛事活动")
    add_bullet(doc, "多租户权限体系、完整审计后台")
    add_bullet(doc, "硬件闸机对接与离线核销强一致方案")
    add_heading_cn(doc, "12.3 后续增强建议", 2)
    add_bullet(doc, "天气/鱼汛信息叠加地图")
    add_bullet(doc, "推荐码分销结算与渠道报表")
    add_bullet(doc, "打卡防作弊（轨迹、围栏、频控）增强")
    add_bullet(doc, "运营活动配置台与消息推送")

    # ========== 附录 ==========
    add_heading_cn(doc, "附录 A：功能清单总表", 1)
    add_table(
        doc,
        ["端", "模块", "核心能力"],
        [
            ["用户端", "地图", "已开通/维护中区域；规则与鱼种"],
            ["用户端", "购票", "年票（推荐码 FISH2026→¥688）、次票 ¥68"],
            ["用户端", "打卡", "GPS 选区确认；位置上传；钓场实时可见"],
            ["用户端", "我的", "年票剩余天数、次票状态、积分、等级、每日任务"],
            ["钓场端", "区域地图", "仅本场；已打卡点位脉冲动画"],
            ["钓场端", "票务核销", "扫码核销年票/次票；核销记录"],
            ["钓场端", "打卡点位", "当前用户列表与打卡时间"],
            ["钓场端", "数据统计", "今日收入、客流、年票/次票持有数"],
            ["后台端", "核心指标", "总用户、钓场数、今日打卡、今日收入"],
            ["后台端", "收入趋势", "周收入柱状图；本月/今日对比"],
            ["后台端", "收益明细", "各场年票/次票拆分、打卡次数、收益占比条"],
        ],
        [3, 3.5, 9.5],
    )

    add_heading_cn(doc, "附录 B：原型操作速查", 1)
    add_para(
        doc,
        "顶部切换角色 → 底部切换功能页。完整路径：用户地图了解规则 → 购票（年票试 FISH2026）→ 打卡上传 → 切到钓场看脉冲点位与核销 → 切到后台核对应指标与收益明细。",
        first_line=True,
    )

    out = "/workspace/docs/垂钓管理平台-产品需求规格说明书-V1.0.docx"
    doc.save(out)
    print(out)


if __name__ == "__main__":
    build()
